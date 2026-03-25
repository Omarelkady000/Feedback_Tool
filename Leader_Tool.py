import streamlit as st
import re
import requests
import math
import csv
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- 1. CORE LOGIC ---
XML_TIMEBASE_MAP = {
    "10.00 fps": "10", "12.00 fps": "12", "15.00 fps": "15",
    "23.976 fps": "24", "24.00 fps": "24", "25.00 fps": "25", "29.97 fps": "30",
    "30.00 fps": "30", "50.00 fps": "50", "59.94 fps": "60", "60.00 fps": "60"
}

def tc_to_frames(tc, fps_choice):
    try:
        clean_tc = tc.replace(';', ':')
        parts = list(map(int, clean_tc.split(':')))
        h, m, s, f = parts
        total_minutes = (h * 60) + m
        if "29.97" in fps_choice:
            frame_number = ((total_minutes * 60) + s) * 30 + f
            drop_frames = 2 * (total_minutes - (total_minutes // 10))
            return frame_number - drop_frames
        elif "59.94" in fps_choice:
            frame_number = ((total_minutes * 60) + s) * 60 + f
            drop_frames = 4 * (total_minutes - (total_minutes // 10))
            return frame_number - drop_frames
        else:
            base = 24 if "23.976" in fps_choice else float(fps_choice.split(' ')[0])
            return math.floor((h * 3600 * base) + (m * 60 * base) + (s * base) + f)
    except: return 0

def apply_style(run, size=11, bold=False, color=None):
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

# --- 2. UI CONFIGURATION ---
st.set_page_config(page_title="Editor Workflow Tool", page_icon="🎬", layout="centered")

st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .stButton>button { width: 100%; border-radius: 5px; height: 3em; background-color: #4CAF50; color: white; }
    </style>
    """, unsafe_allow_html=True)

st.title("🎬 Premiere Workflow Suite")
st.info("Select your frame rate in the sidebar before processing.")

# Sidebar Settings
st.sidebar.header("Global Settings")
fps_choice = st.sidebar.selectbox("Sequence FPS:", list(XML_TIMEBASE_MAP.keys()), index=6)

tab1, tab2 = st.tabs(["📤 Editor: Doc to XML", "📥 Leader: CSV to Docx/XML"])

# --- TAB 1: EDITOR MODE ---
with tab1:
    st.subheader("Convert Google Doc Comments to Markers")
    url = st.text_input("Google Doc URL:", placeholder="https://docs.google.com/document/d/...")
    custom_name = st.text_input("Output Filename:", value="Feedback_Markers")

    if st.button("Generate Premiere XML", key="btn_xml"):
        if not url:
            st.warning("Please provide a Google Doc URL.")
        else:
            try:
                export_url = url.split('/edit')[0] + '/export?format=txt' if "/edit" in url else url
                response = requests.get(export_url)
                data = response.text
                pattern = r"(\d{2}[:;]\d{2}[:;]\d{2}[:;]\d{2})(?:\s*[–-]\s*(\d{2}[:;]\d{2}[:;]\d{2}[:;]\d{2}))?([\s\S]+?)(?=\d{2}[:;]\d{2}[:;]\d{2}[:;]\d{2}|$)"
                matches = list(re.finditer(pattern, data))

                timebase = XML_TIMEBASE_MAP.get(fps_choice, "30")
                ntsc = "TRUE" if (".97" in fps_choice or ".94" in fps_choice) else "FALSE"
                
                xml = f'<?xml version="1.0" encoding="UTF-8"?><xmeml version="4"><project><children><sequence><name>FEEDBACK</name><rate><timebase>{timebase}</timebase><ntsc>{ntsc}</ntsc></rate><media><video><format><samplecharacteristics><width>1920</width><height>1080</height></samplecharacteristics></format></video></media>'
                
                for m in matches:
                    start_f = tc_to_frames(m.group(1), fps_choice)
                    end_f = tc_to_frames(m.group(2), fps_choice) if m.group(2) else start_f
                    cmt = m.group(3).strip()
                    if bool(re.search(r'\bkeep\b', cmt, re.IGNORECASE)): end_f = start_f
                    clean_cmt = cmt.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    xml += f"<marker><name>NOTE</name><comment>{clean_cmt}</comment><in>{int(start_f)}</in><out>{int(end_f)}</out></marker>"
                
                xml += "</sequence></children></project></xmeml>"
                st.download_button("⬇️ Download XML File", data=xml, file_name=f"{custom_name}.xml", mime="application/xml")
            except Exception as e:
                st.error(f"Failed to fetch document: {e}")

# --- TAB 2: LEADER MODE ---
with tab2:
    st.subheader("Convert Premiere CSV to Formatted Feedback")
    csv_file = st.file_uploader("Upload Markers CSV", type="csv")
    logo_file = st.file_uploader("Upload Brand Logo (Optional)", type=["png", "jpg"])

    if csv_file:
        try:
            raw_data = csv_file.read()
            content = ""
            for enc in ['utf-8-sig', 'utf-16', 'cp1252']:
                try:
                    content = raw_data.decode(enc)
                    if "Marker Name" in content: break
                except: continue
            
            dialect = csv.Sniffer().sniff(content[:2000])
            reader = csv.DictReader(content.splitlines(), dialect=dialect)
            
            doc = Document()
            if logo_file:
                doc.add_picture(io.BytesIO(logo_file.read()), width=Inches(1.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Title
            title_name = Path(csv_file.name).stem
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_h = header.add_run(title_name)
            apply_style(run_h, size=16, bold=True)

            xml_markers = ""
            for row in reader:
                name = row.get('Marker Name', '').strip()
                desc = row.get('Description', '').strip()
                comment = name if len(name) >= len(desc) else desc
                in_tc = row.get('In', '00:00:00:00')
                out_tc = row.get('Out', in_tc)
                
                ts_display = in_tc if in_tc == out_tc else f"{in_tc} - {out_tc}"
                is_neg = bool(re.search(r'\b(remove|cut|delete)\b', comment, re.IGNORECASE))
                text_color = RGBColor(255, 0, 0) if is_neg else None
                
                p = doc.add_paragraph()
                run_ts = p.add_run(f"{ts_display}  ")
                apply_style(run_ts, bold=True, color=text_color)
                run_cmt = p.add_run(comment)
                apply_style(run_cmt, color=text_color)
                
                # XML data
                start_f = tc_to_frames(in_tc, fps_choice)
                end_f = tc_to_frames(out_tc, fps_choice)
                clean_cmt = comment.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                xml_markers += f"<marker><name>NOTE</name><comment>{clean_cmt}</comment><in>{int(start_f)}</in><out>{int(end_f)}</out></marker>"

            # Preparation for download
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            timebase = XML_TIMEBASE_MAP.get(fps_choice, "30")
            ntsc = "TRUE" if (".97" in fps_choice or ".94" in fps_choice) else "FALSE"
            full_xml = f'<?xml version="1.0" encoding="UTF-8"?><xmeml version="4"><project><children><sequence><name>CSV_IMPORT</name><rate><timebase>{timebase}</timebase><ntsc>{ntsc}</ntsc></rate><media><video><format><samplecharacteristics><width>1920</width><height>1080</height></samplecharacteristics></format></video></media>{xml_markers}</sequence></children></project></xmeml>'

            c1, c2 = st.columns(2)
            c1.download_button("⬇️ Download Docx", data=doc_buffer, file_name=f"{title_name}.docx")
            c2.download_button("⬇️ Download XML", data=full_xml, file_name=f"{title_name}.xml")
        except Exception as e:
            st.error(f"Error: {e}")
