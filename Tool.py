import streamlit as st
import re
import math
import csv
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- 1. CONFIGURATION DATA ---
XML_TIMEBASE_MAP = {
    "10.00 fps": "10", "12.00 fps": "12", "15.00 fps": "15",
    "23.976 fps": "24", "24.00 fps": "24", "25.00 fps": "25", "29.97 fps": "30",
    "30.00 fps": "30", "50.00 fps": "50", "59.94 fps": "60", "60.00 fps": "60"
}

# Added Resolution Mapping
RES_MAP = {
    "1080x1920 (Vertical HD)": (1080, 1920),
    "1920x1080 (Landscape HD)": (1920, 1080),
    "2160x3840 (Vertical 4K)": (2160, 3840),
    "3840x2160 (Landscape 4K)": (3840, 2160),
    "1080x1080 (Square)": (1080, 1080)
}

def tc_to_frames(tc, fps_choice):
    try:
        clean_tc = tc.replace(';', ':')
        parts = list(map(int, clean_tc.split(':')))
        h, m, s, f = parts
        total_minutes = (h * 60) + m
        if "29.97" in fps_choice:
            frame_number = ((total_minutes * 60) + s) * 30 + f
            drop_frames = 2 * (total_minutes - (total_minutes // 10))
            return frame_number - drop_frames
        elif "59.94" in fps_choice:
            frame_number = ((total_minutes * 60) + s) * 60 + f
            drop_frames = 4 * (total_minutes - (total_minutes // 10))
            return frame_number - drop_frames
        else:
            base = 24 if "23.976" in fps_choice else float(fps_choice.split(' ')[0])
            return math.floor((h * 3600 * base) + (m * 60 * base) + (s * base) + f)
    except: return 0

def set_font(run, size=11, bold=False):
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    run.font.size = Pt(size)
    run.bold = bold

# --- 2. UI SETUP ---
st.set_page_config(page_title="QOMY Feedback Tool", page_icon="🎬", layout="centered")

st.title("🎬 QOMY Feedback Tool")
st.markdown("Upload your Premiere CSV to generate formatted Feedback Docs and XML Markers.")

# Sidebar Settings
st.sidebar.header("GLOBAL SETTINGS")

st.sidebar.write("Select Premiere Sequence FPS:")
fps_choice = st.sidebar.selectbox("FPS Dropdown:", list(XML_TIMEBASE_MAP.keys()), index=6, label_visibility="collapsed")

st.sidebar.write("Select Sequence Resolution:")
res_choice = st.sidebar.selectbox("Resolution Dropdown:", list(RES_MAP.keys()), index=0, label_visibility="collapsed")
width, height = RES_MAP[res_choice]

# --- 3. WORKFLOW ---
csv_file = st.file_uploader("Select Premiere CSV", type="csv")
logo_file = st.file_uploader("Upload Logo (Optional)", type=["png", "jpg"])

if csv_file:
    try:
        raw_data = csv_file.read()
        content = ""
        for enc in ['utf-8-sig', 'utf-16', 'cp1252']:
            try:
                content = raw_data.decode(enc)
                if "Marker Name" in content: break
            except: continue
        
        dialect = csv.Sniffer().sniff(content[:2000])
        reader = csv.DictReader(content.splitlines(), dialect=dialect)
        
        doc = Document()
        
        # 1. Logo
        if logo_file:
            doc.add_picture(io.BytesIO(logo_file.read()), width=Inches(1.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 2. Title Logic
        base_name = Path(csv_file.name).stem
        final_filename = f"{base_name}feedback" # Adjusted per request 1
        
        title_para = doc.add_heading('', 0)
        title_run = title_para.add_run(base_name)
        set_font(title_run, size=18, bold=True)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        xml_markers = ""
        
        for row in reader:
            name = row.get('Marker Name', '').strip()
            desc = row.get('Description', '').strip()
            comment = name if len(name) >= len(desc) else desc
            
            in_tc = row.get('In', '00:00:00:00')
            out_tc = row.get('Out', in_tc)
            ts_display = in_tc if in_tc == out_tc else f"{in_tc} - {out_tc}"
            
            # Word Doc Paragraphs
            p_ts = doc.add_paragraph()
            run_ts = p_ts.add_run(ts_display)
            set_font(run_ts, bold=True)

            p_cmt = doc.add_paragraph()
            run_cmt = p_cmt.add_run(comment)
            set_font(run_cmt)
            
            # XML Logic
            start_f = tc_to_frames(in_tc, fps_choice)
            end_f = tc_to_frames(out_tc, fps_choice)
            clean_cmt = comment.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            xml_markers += f"<marker><name>NOTE</name><comment>{clean_cmt}</comment><in>{int(start_f)}</in><out>{int(end_f)}</out></marker>"

        # Prepare Buffers
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        timebase = XML_TIMEBASE_MAP.get(fps_choice, "30")
        ntsc = "TRUE" if (".97" in fps_choice or ".94" in fps_choice) else "FALSE"
        
        # Adjusted Sequence Name per request 1 and Resolution per request 2
        full_xml = f'<?xml version="1.0" encoding="UTF-8"?><xmeml version="4"><project><children><sequence><name>{final_filename}</name><rate><timebase>{timebase}</timebase><ntsc>{ntsc}</ntsc></rate><media><video><format><samplecharacteristics><width>{width}</width><height>{height}</height></samplecharacteristics></format></video></media>{xml_markers}</sequence></children></project></xmeml>'

        st.divider()
        st.success(f"Processed: {base_name}")
        
        c1, c2 = st.columns(2)
        with c1:
            st.download_button("⬇️ Download Docx", data=doc_io, file_name=f"{final_filename}.docx")
        with c2:
            st.download_button("⬇️ Download XML", data=full_xml, file_name=f"{final_filename}.xml")

    except Exception as e:
        st.error(f"Error: {e}")
